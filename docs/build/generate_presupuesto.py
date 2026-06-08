"""
Genera AgroCore - Presupuesto.docx con estilo profesional.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
from pathlib import Path
import random

BASE = Path("/sessions/amazing-trusting-einstein/mnt/AgroCore")
LOGO_FULL = BASE / "web/img/logo-full-512.png"
LOGO_ICON = BASE / "web/img/logo-icon-128.png"
OUT = BASE / "docs/Presupuesto-AgroCore.docx"

AGRO_DARK = RGBColor(0x14, 0x53, 0x2D)
AGRO_GREEN = RGBColor(0x15, 0x80, 0x3D)
AGRO_MED = RGBColor(0x16, 0x65, 0x34)
GOLD = RGBColor(0xCA, 0x8A, 0x04)
GOLD_LIGHT = RGBColor(0xEA, 0xB3, 0x08)
BLUE_DARK = RGBColor(0x1E, 0x40, 0xAF)
AMBER_DARK = RGBColor(0xB4, 0x53, 0x09)
SLATE_700 = RGBColor(0x33, 0x41, 0x55)
SLATE_500 = RGBColor(0x64, 0x74, 0x8B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)

def set_cell_borders(cell, color="CCCCCC", size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), size); b.set(qn('w:color'), color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)

def no_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right'):
        b = OxmlElement(f'w:{edge}'); b.set(qn('w:val'), 'nil')
        tc_borders.append(b)
    tc_pr.append(tc_borders)

def add_page_number(paragraph):
    run = paragraph.add_run()
    f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = 'PAGE'
    f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'end')
    run._r.append(f1); run._r.append(it); run._r.append(f2)

def add_horizontal_line(paragraph, color="22C55E"):
    p_pr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '12'); bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), color)
    pbdr.append(bottom)
    p_pr.append(pbdr)

def h1(doc, text, color=AGRO_DARK):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16); p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name='Calibri'; r.font.size=Pt(20); r.font.bold=True; r.font.color.rgb=color
    add_horizontal_line(p, "DCFCE7")

def h2(doc, text, color=AGRO_GREEN):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name='Calibri'; r.font.size=Pt(14); r.font.bold=True; r.font.color.rgb=color

def h3(doc, text, color=AGRO_MED):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name='Calibri'; r.font.size=Pt(12); r.font.bold=True; r.font.color.rgb=color

def para(doc, text, bold=False, italic=False, color=None, size=11, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if align: p.alignment = align
    r = p.add_run(text)
    r.font.name = 'Calibri'; r.font.size = Pt(size)
    r.font.bold = bold; r.font.italic = italic
    if color: r.font.color.rgb = color
    return p

def para_mixed(doc, parts, size=11, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if align: p.alignment = align
    for text, opts in parts:
        r = p.add_run(text)
        r.font.name = 'Consolas' if opts.get('code') else 'Calibri'
        r.font.size = Pt(size - 1 if opts.get('code') else size)
        r.font.bold = opts.get('bold', False)
        r.font.italic = opts.get('italic', False)
        if opts.get('color'): r.font.color.rgb = opts['color']
    return p

def bullet(doc, text_or_parts, check_color=AGRO_GREEN):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.25)
    if isinstance(text_or_parts, str):
        r = p.add_run(text_or_parts)
        r.font.name='Calibri'; r.font.size=Pt(11)
    else:
        for text, opts in text_or_parts:
            r = p.add_run(text)
            r.font.name = 'Consolas' if opts.get('code') else 'Calibri'
            r.font.size = Pt(10 if opts.get('code') else 11)
            r.font.bold = opts.get('bold', False)

# ============================================================
# Build document
# ============================================================
doc = Document()

for section in doc.sections:
    section.page_height = Cm(29.7); section.page_width = Cm(21.0)
    section.top_margin = Cm(2.0); section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0); section.right_margin = Cm(2.0)

style = doc.styles['Normal']
style.font.name = 'Calibri'; style.font.size = Pt(11); style.font.color.rgb = SLATE_700

# Footer
footer = doc.sections[0].footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = fp.add_run("AgroCore Argentina · Presupuesto · Página ")
r.font.name='Calibri'; r.font.size=Pt(9); r.font.color.rgb=SLATE_500
add_page_number(fp)

# Header
header = doc.sections[0].header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = hp.add_run("agrocore.ar")
r.font.name='Calibri'; r.font.size=Pt(9); r.font.color.rgb=AGRO_GREEN; r.font.bold=True

# ============================================================
# COVER PAGE
# ============================================================
cover_p = doc.add_paragraph()
cover_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
# Cover decoration
    _cd = doc.add_paragraph(); _cd.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _cd.paragraph_format.space_before = Pt(0); _cd.paragraph_format.space_after = Pt(15)
    _cd.add_run().add_picture(str(Path("/sessions/amazing-trusting-einstein/mnt/AgroCore/docs/build/imgs/cover_decoration.png")), width=Inches(7))
    cover_p.paragraph_format.space_before = Pt(10)
if LOGO_FULL.exists():
    cover_p.add_run().add_picture(str(LOGO_FULL), width=Inches(3))

# Badge "Presupuesto"
badge_p = doc.add_paragraph(); badge_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
badge_p.paragraph_format.space_before = Pt(20)
r = badge_p.add_run("PRESUPUESTO")
r.font.name='Calibri'; r.font.size=Pt(12); r.font.bold=True; r.font.color.rgb=GOLD

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.paragraph_format.space_before = Pt(6)
r = t.add_run("Sistema integral de gestión rural")
r.font.name='Calibri'; r.font.size=Pt(28); r.font.bold=True; r.font.color.rgb=AGRO_DARK

sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Desarrollo, implementación, capacitación y soporte.")
r.font.name='Calibri'; r.font.size=Pt(13); r.font.color.rgb=SLATE_700
sub2 = doc.add_paragraph(); sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run("Multi-empresa, multi-usuario, en la nube o en servidor propio.")
r.font.name='Calibri'; r.font.size=Pt(13); r.font.color.rgb=SLATE_700

# Document metadata box
hoy = datetime.now()
nro = f"AC-{hoy.year}-{random.randint(1000, 9999)}"
fecha_str = hoy.strftime("%d/%m/%Y")

meta_t = doc.add_table(rows=1, cols=2)
meta_t.autofit = False
meta_t.columns[0].width = Inches(3.25); meta_t.columns[1].width = Inches(3.25)
meta_t.alignment = WD_ALIGN_PARAGRAPH.CENTER

c1 = meta_t.cell(0, 0); c2 = meta_t.cell(0, 1)
c1.width = Inches(3.25); c2.width = Inches(3.25)
no_cell_borders(c1); no_cell_borders(c2)
set_cell_shading(c1, "F0FDF4"); set_cell_shading(c2, "F0FDF4")

p = c1.paragraphs[0]
p.paragraph_format.space_before = Pt(60)
r = p.add_run("EMITIDO\n"); r.font.size=Pt(8); r.font.color.rgb=SLATE_500; r.font.bold=True
r = c1.add_paragraph().add_run(fecha_str); r.font.size=Pt(13); r.font.bold=True; r.font.color.rgb=AGRO_DARK

p = c2.paragraphs[0]
p.paragraph_format.space_before = Pt(60)
r = p.add_run("N°\n"); r.font.size=Pt(8); r.font.color.rgb=SLATE_500; r.font.bold=True
r = c2.add_paragraph().add_run(nro); r.font.size=Pt(13); r.font.bold=True; r.font.color.rgb=AGRO_DARK

validez = doc.add_paragraph(); validez.alignment = WD_ALIGN_PARAGRAPH.CENTER
validez.paragraph_format.space_before = Pt(20)
r = validez.add_run("Válido por 30 días desde la fecha de emisión")
r.font.name='Calibri'; r.font.size=Pt(10); r.font.italic=True; r.font.color.rgb=SLATE_500

doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

# ============================================================
# RESUMEN EJECUTIVO
# ============================================================
h1(doc, "Resumen ejecutivo")
para_mixed(doc, [
    ("AgroCore es un sistema integral de gestión ", {}),
    ("diseñado específicamente para el agro argentino", {"bold": True, "color": AGRO_DARK}),
    (". Reemplaza planillas dispersas, sistemas antiguos y procesos manuales por una única plataforma que integra producción, stock, comercial, tesorería, hacienda y empleados.", {}),
])
para_mixed(doc, [
    ("Esta propuesta detalla la ", {}),
    ("inversión única de desarrollo + implementación", {"bold": True}),
    (", junto con dos componentes opcionales: ", {}),
    ("mantenimiento mensual", {"bold": True}),
    (" e ", {}),
    ("infraestructura", {"bold": True}),
    (" (a elección entre Cloud sobre AWS o servidor propio en tu empresa).", {}),
])

# 3-card preview of investment
inv_t = doc.add_table(rows=1, cols=3)
inv_t.autofit = False
inv_t.columns[0].width = Inches(2.15); inv_t.columns[1].width = Inches(2.15); inv_t.columns[2].width = Inches(2.15)
preview = [
    ("F0FDF4", "166534", "USD 999,99", "Pago único · Desarrollo + Implementación"),
    ("FEF3C7", "B45309", "USD 100/mes", "Mantenimiento opcional"),
    ("DBEAFE", "1E40AF", "A elección", "Cloud (USD 20–40/mes) o Servidor propio (sin costo)"),
]
for i, (bg, fg, amount, desc) in enumerate(preview):
    c = inv_t.cell(0, i); c.width = Inches(2.15)
    set_cell_shading(c, bg)
    set_cell_borders(c, fg, "8")
    c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run(amount); r.font.name='Calibri'; r.font.size=Pt(16); r.font.bold=True; r.font.color.rgb=RGBColor.from_string(fg)
    p2 = c.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(8)
    r = p2.add_run(desc); r.font.name='Calibri'; r.font.size=Pt(9); r.font.color.rgb=SLATE_700

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ============================================================
# ¿QUÉ INCLUYE?
# ============================================================
h1(doc, "¿Qué incluye el sistema?")
para(doc, "7 secciones con más de 24 vistas integradas. Sin costos por usuario adicional ni por funcionalidad extra dentro del alcance.")

modulos = [
    ("🏠 Inicio + Dashboard",
     ["KPIs en tiempo real", "Cotizaciones BCR en vivo",
      "5 tabs (Resumen, Finanzas, Stock, Producción, Patrimonial)",
      "Gráficos y comparativas",
      "Resumen multi-empresa consolidado"]),
    ("🌱 Producción",
     ["Campos y lotes con ambientes",
      "Campañas (gruesa, fina, doble cultivo)",
      "Insumos y labores aplicadas (con hectáreas custom)",
      "Margen bruto en tiempo real"]),
    ("📦 Stock",
     ["Inventario valorizado",
      "Movimientos con trazabilidad",
      "Reporte de margen (compras vs ventas)",
      "Alertas de stock mínimo"]),
    ("👥 Contactos",
     ["Clientes (con cuenta corriente)",
      "Proveedores (con cuenta corriente)",
      "Saldo y total facturado en vivo",
      "Persona de contacto y notas internas"]),
    ("💰 Finanzas (la sección más amplia)",
     ["Flujo de caja proyectado 30/60/90 días",
      "Facturación electrónica ARCA (CAE)",
      "Compras (con importación masiva)",
      "Libro IVA Ventas y Compras",
      "Cheques (ciclo completo con alertas)",
      "Cuentas corrientes con conciliación",
      "Arrendamientos en quintales / USD / pesos",
      "Control de efectivo (caja chica)"]),
    ("🚚 Logística y RRHH",
     ["Viajes con cartas de porte (CTG)",
      "Empleados con contratos",
      "Recibos de sueldo automáticos",
      "Aguinaldo, vacaciones, horas extras"]),
    ("⚙️ Administración",
     ["Catálogos (productos, cultivos, bancos)",
      "Usuarios con asignación a empresas",
      "Roles y permisos a medida",
      "Empresas con configuración ARCA"]),
]
for titulo, items in modulos:
    h3(doc, titulo)
    for it in items:
        bullet(doc, it)

# Bonus box
bonus_t = doc.add_table(rows=1, cols=1)
bonus_t.autofit = False
bonus_t.columns[0].width = Inches(6.5)
c = bonus_t.cell(0, 0); c.width = Inches(6.5)
set_cell_shading(c, "F0FDF4")
set_cell_borders(c, "BBF7D0", "6")
p = c.paragraphs[0]
p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
r = p.add_run("Bonus: ")
r.font.bold = True; r.font.color.rgb = AGRO_GREEN
r = p.add_run("Multi-empresa sin costo extra · Multi-usuario sin tope · RBAC con permisos granulares · Auditoría completa de acciones · PWA con modo offline · Navegación con botón Volver en toda la app.")
r.font.size = Pt(10.5); r.font.color.rgb = SLATE_700

doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

# ============================================================
# DETALLE DE LA INVERSIÓN
# ============================================================
h1(doc, "Detalle de la inversión")
para_mixed(doc, [
    ("Todos los precios en dólares estadounidenses (USD). Pago en pesos al tipo de cambio Banco Nación + 1% del día de pago, salvo acuerdo.", {"italic": True}),
])

# Card 1: Desarrollo
def add_concept_card(doc, badge_text, badge_color, title, amount, amount_color, subtitle, items_with_check, check_color=GOLD, bg="14532D", txt_white=True):
    t = doc.add_table(rows=2, cols=2)
    t.autofit = False
    t.columns[0].width = Inches(4.25); t.columns[1].width = Inches(2.25)
    # Top-left: badge + title
    c00 = t.cell(0, 0); c00.width = Inches(4.25)
    set_cell_shading(c00, bg); no_cell_borders(c00)
    p = c00.paragraphs[0]; p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(badge_text); r.font.size=Pt(8); r.font.bold=True
    r.font.color.rgb = badge_color if not txt_white else GOLD_LIGHT
    p2 = c00.add_paragraph()
    p2.paragraph_format.space_after = Pt(8)
    r = p2.add_run(title)
    r.font.size = Pt(18); r.font.bold = True
    r.font.color.rgb = WHITE if txt_white else AGRO_DARK
    # Top-right: amount
    c01 = t.cell(0, 1); c01.width = Inches(2.25)
    set_cell_shading(c01, bg); no_cell_borders(c01)
    c01.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = c01.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(0)
    r = p.add_run(amount); r.font.size=Pt(20); r.font.bold=True; r.font.color.rgb=amount_color
    p2 = c01.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.paragraph_format.space_after = Pt(8)
    r = p2.add_run(subtitle); r.font.size=Pt(8)
    r.font.color.rgb = WHITE if txt_white else SLATE_500
    # Bottom merged row: items list
    c10 = t.cell(1, 0); c11 = t.cell(1, 1)
    c10.merge(c11)
    cmerge = t.cell(1, 0)
    set_cell_shading(cmerge, bg); no_cell_borders(cmerge)
    cmerge.paragraphs[0].paragraph_format.space_before = Pt(0)
    first = True
    for it in items_with_check:
        if first:
            p = cmerge.paragraphs[0]
            first = False
        else:
            p = cmerge.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run("✓  "); r.font.size=Pt(11); r.font.bold=True; r.font.color.rgb=check_color
        r = p.add_run(it); r.font.size=Pt(10)
        r.font.color.rgb = WHITE if txt_white else SLATE_700
    # Padding at bottom
    p = cmerge.add_paragraph(); p.paragraph_format.space_after = Pt(8)

add_concept_card(doc,
    badge_text="CONCEPTO 1 · PAGO ÚNICO",
    badge_color=GOLD_LIGHT,
    title="Desarrollo + Implementación",
    amount="USD 999,99",
    amount_color=GOLD_LIGHT,
    subtitle="al finalizar la implementación",
    items_with_check=[
        "Sistema completo con las 7 secciones (24+ vistas integradas)",
        "Implementación en infraestructura a elección (Cloud AWS o Local)",
        "Carga inicial de catálogos maestros (productos, cultivos, bancos)",
        "Migración de datos desde Excel o sistema previo (volumen razonable)",
        "Capacitación remota inicial (3 sesiones de 1 hora vía Zoom/Meet)",
        "Manual de usuario online + documentación de arquitectura",
        "URL pública con HTTPS y SSL gratuito",
        "Configuración de integración con ARCA (certificado y punto de venta)",
        "Garantía de 30 días post-implementación (corrección de bugs sin costo)",
        "Acuerdo de Confidencialidad (NDA) firmado",
    ],
    bg="14532D")

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# Card 2: Mantenimiento
add_concept_card(doc,
    badge_text="CONCEPTO 2 · OPCIONAL · CANCELABLE",
    badge_color=AMBER_DARK,
    title="Mantenimiento Mensual",
    amount="USD 100/mes",
    amount_color=AGRO_DARK,
    subtitle="cancelable cuando quieras",
    items_with_check=[
        "Soporte técnico (WhatsApp y email, respuesta en 24 hs hábiles)",
        "Corrección de bugs sin costo",
        "Actualizaciones legales ARCA (cambios de alícuotas, nuevos comprobantes)",
        "Resguardo y verificación de backups",
        "Monitoreo del servidor 24/7 (con alerta proactiva)",
        "3 hs/mes de ajustes menores (campos extra, reportes a medida)",
        "Actualización de cotizaciones BCR",
        "Acceso a nuevas features publicadas en versiones menores",
    ],
    check_color=AGRO_GREEN,
    bg="FEF3C7",
    txt_white=False)

# Note about mantenimiento
note_t = doc.add_table(rows=1, cols=1)
note_t.autofit = False
note_t.columns[0].width = Inches(6.5)
c = note_t.cell(0, 0); c.width = Inches(6.5)
set_cell_shading(c, "FEF3C7"); set_cell_borders(c, "FCD34D", "4")
p = c.paragraphs[0]; p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
r = p.add_run('"Lo tomás, lo dejás, lo retomás": ')
r.font.bold = True; r.font.color.rgb = AMBER_DARK
r = p.add_run("el mantenimiento es 100% opcional. Si lo cancelás, el sistema sigue funcionando perfectamente. Si más adelante lo necesitás de nuevo, podés re-suscribirte sin penalidades.")
r.font.size = Pt(10); r.font.color.rgb = SLATE_700

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# Card 3: Infraestructura - two options
h2(doc, "Concepto 3 · A elección del cliente · Infraestructura")
para(doc, "Elegís dónde corre el sistema. En ambos casos los datos y el código son tuyos.")

# Cloud option
add_concept_card(doc,
    badge_text="OPCIÓN A · NUBE",
    badge_color=BLUE_DARK,
    title="☁️ Cloud (AWS)",
    amount="USD 20–40/mes",
    amount_color=BLUE_DARK,
    subtitle="estimado · pago directo a Amazon",
    items_with_check=[
        "Cuenta AWS a tu nombre",
        "Backups automáticos diarios (RDS)",
        "Acceso desde cualquier lugar con internet",
        "Encriptación AES-256, certificado SSL",
        "Alertas de presupuesto AWS",
        "Costo escala según uso real",
    ],
    check_color=BLUE_DARK,
    bg="DBEAFE",
    txt_white=False)

# Recommendation note
note_t = doc.add_table(rows=1, cols=1); note_t.autofit = False
note_t.columns[0].width = Inches(6.5)
c = note_t.cell(0, 0); c.width = Inches(6.5)
set_cell_shading(c, "DBEAFE"); set_cell_borders(c, "BFDBFE", "4")
p = c.paragraphs[0]; p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
r = p.add_run("Recomendado para empresas con conectividad estable y múltiples sucursales o usuarios en distintos lugares.")
r.font.size = Pt(9); r.font.italic = True; r.font.color.rgb = BLUE_DARK
doc.add_paragraph().paragraph_format.space_after = Pt(6)

# On-premise option
add_concept_card(doc,
    badge_text="OPCIÓN B · ON-PREMISE",
    badge_color=AMBER_DARK,
    title="🖥️ Servidor Propio",
    amount="Sin costo recurrente",
    amount_color=AMBER_DARK,
    subtitle="instalado en tu empresa",
    items_with_check=[
        "Instalación on-premise en tu servidor",
        "Datos siempre dentro de tu oficina/galpón",
        "Funciona sin internet (en LAN)",
        "Acceso remoto vía VPN (opcional)",
        "Backups a tu cargo (te ayudamos a configurarlos)",
        "Hardware mínimo: 8 GB RAM, 250 GB disco, Windows o Linux",
        "Sin facturas mensuales de cloud",
    ],
    check_color=AMBER_DARK,
    bg="FEF3C7",
    txt_white=False)

note_t = doc.add_table(rows=1, cols=1); note_t.autofit = False
note_t.columns[0].width = Inches(6.5)
c = note_t.cell(0, 0); c.width = Inches(6.5)
set_cell_shading(c, "FEF3C7"); set_cell_borders(c, "FCD34D", "4")
p = c.paragraphs[0]; p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
r = p.add_run("Recomendado para empresas con un servidor disponible y preferencia por mantener los datos físicamente en sus instalaciones.")
r.font.size = Pt(9); r.font.italic = True; r.font.color.rgb = AMBER_DARK

# Investment summary in dark box
doc.add_paragraph().paragraph_format.space_after = Pt(10)

sum_t = doc.add_table(rows=1, cols=1); sum_t.autofit = False
sum_t.columns[0].width = Inches(6.5)
c = sum_t.cell(0, 0); c.width = Inches(6.5)
set_cell_shading(c, "14532D"); set_cell_borders(c, "0F3D20", "4")

p = c.paragraphs[0]; p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(4)
r = p.add_run("Inversión total estimada")
r.font.size = Pt(14); r.font.bold = True; r.font.color.rgb = GOLD_LIGHT

for label, value in [
    ("Desarrollo + Implementación (pago único)", "USD 999,99"),
    ("Infraestructura mensual (opción Cloud)", "USD 20–40 / mes"),
    ("Infraestructura mensual (opción Servidor Propio)", "Sin costo"),
    ("Mantenimiento mensual (opcional)", "USD 100 / mes"),
]:
    p = c.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label); r.font.size=Pt(10); r.font.color.rgb=WHITE
    r = p.add_run("\t"); r.font.color.rgb=WHITE
    r = p.add_run(value); r.font.size=Pt(10); r.font.bold=True; r.font.color.rgb=WHITE

p = c.add_paragraph(); p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
r = p.add_run("Ejemplo cliente A: "); r.font.bold = True; r.font.size = Pt(9); r.font.color.rgb = GOLD_LIGHT
r = p.add_run("Cloud + mantenimiento → primer pago USD 999,99 (única vez), luego ~USD 130/mes recurrente.")
r.font.size = Pt(9); r.font.color.rgb = WHITE
p = c.add_paragraph(); p.paragraph_format.space_after = Pt(8)
r = p.add_run("Ejemplo cliente B: "); r.font.bold = True; r.font.size = Pt(9); r.font.color.rgb = GOLD_LIGHT
r = p.add_run("Servidor propio sin mantenimiento → primer pago USD 999,99 (única vez), luego USD 0/mes recurrente.")
r.font.size = Pt(9); r.font.color.rgb = WHITE

doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

# ============================================================
# CRONOGRAMA
# ============================================================
h1(doc, "Cronograma de implementación")
para(doc, "Entre 4 y 6 semanas desde la confirmación, dependiendo del volumen de datos a migrar.")

cron_t = doc.add_table(rows=2, cols=4); cron_t.autofit = False
cw = 1.625
for i in range(4): cron_t.columns[i].width = Inches(cw)

fases = [
    ("1", "Charlamos", "Semana 1\nEntendemos tu operatoria"),
    ("2", "Configuramos", "Semanas 2-3\nInfraestructura + carga inicial"),
    ("3", "Capacitamos", "Semana 4\n3 sesiones con tu equipo"),
    ("4", "Salimos en vivo", "Semanas 5-6\nGarantía 30 días incluida"),
]
for i, (num, titulo, desc) in enumerate(fases):
    # Top row: numbered circle + title
    c = cron_t.cell(0, i); c.width = Inches(cw)
    set_cell_shading(c, "F0FDF4"); set_cell_borders(c, "BBF7D0", "6")
    c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(num); r.font.size=Pt(22); r.font.bold=True; r.font.color.rgb=AGRO_GREEN
    p2 = c.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(6)
    r = p2.add_run(titulo); r.font.size=Pt(11); r.font.bold=True; r.font.color.rgb=AGRO_DARK
    # Bottom row: description
    c2 = cron_t.cell(1, i); c2.width = Inches(cw)
    set_cell_shading(c2, "F0FDF4"); set_cell_borders(c2, "BBF7D0", "6")
    p = c2.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(8)
    r = p.add_run(desc); r.font.size=Pt(9); r.font.color.rgb=SLATE_700

doc.add_paragraph().paragraph_format.space_after = Pt(10)

# ============================================================
# TÉRMINOS Y CONDICIONES
# ============================================================
h1(doc, "Términos y condiciones")

terms = [
    ("Forma de pago — Desarrollo + Implementación",
     "50% al inicio del proyecto (al firmar la propuesta). 50% al finalizar la implementación, una vez validado el sistema en uso real."),
    ("Forma de pago — Mantenimiento mensual",
     "Facturación mensual por adelantado, primer día hábil del mes. Pago vía transferencia bancaria."),
    ("Tipo de cambio",
     "Pago en pesos argentinos al tipo de cambio Banco Nación vendedor + 1% del día de pago, salvo acuerdo distinto."),
    ("Validez de la oferta",
     "Esta propuesta tiene una validez de 30 días corridos desde la fecha de emisión. Pasado ese plazo, los valores y plazos pueden ser revisados."),
    ("Confidencialidad",
     "Toda información del cliente está protegida por NDA mutuo. Los datos cargados en el sistema son propiedad exclusiva del cliente y nunca son utilizados con otro fin."),
    ("Propiedad del sistema",
     "El cliente recibe el código fuente y la infraestructura (cuenta AWS o servidor propio) a su nombre. Sin lock-in de proveedor: podés cambiar de implementador sin perder datos ni código."),
    ("Soporte fuera de horario",
     "El plan de mantenimiento incluye horario L-V 9 a 18 hs. Soporte 24/7 disponible bajo presupuesto adicional."),
    ("Funcionalidades fuera de alcance",
     "Pedidos de módulos adicionales o integraciones específicas (con sistemas de terceros, IoT, balanzas, etc.) se cotizan por separado tras evaluación."),
]
for titulo, contenido in terms:
    h3(doc, titulo)
    para(doc, contenido)

doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

# ============================================================
# CTA / Contacto
# ============================================================
cta_t = doc.add_table(rows=1, cols=1); cta_t.autofit = False
cta_t.columns[0].width = Inches(6.5)
c = cta_t.cell(0, 0); c.width = Inches(6.5)
set_cell_shading(c, "166534"); set_cell_borders(c, "14532D", "8")

p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20); p.paragraph_format.space_after = Pt(8)
r = p.add_run("¿Hablamos?"); r.font.size=Pt(28); r.font.bold=True; r.font.color.rgb=WHITE

p = c.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(20)
r = p.add_run("Una primera charla de 30 minutos sin compromiso. Te mostramos el sistema, vemos si encaja con tu operatoria, y respondemos todas tus preguntas.")
r.font.size=Pt(11); r.font.color.rgb=WHITE

p = c.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(8)
r = p.add_run("📧 sergiodbilbao@gmail.com"); r.font.size=Pt(12); r.font.bold=True; r.font.color.rgb=GOLD_LIGHT

p = c.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
r = p.add_run("🌐 agrocore.ar"); r.font.size=Pt(12); r.font.color.rgb=WHITE

p = c.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(20)
r = p.add_run("🎯 Probar el demo: demo.agrocore.ar"); r.font.size=Pt(11); r.font.color.rgb=WHITE

p = c.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(16)
r = p.add_run("Río Cuarto, Córdoba · Argentina"); r.font.size=Pt(9); r.font.italic=True; r.font.color.rgb=GOLD_LIGHT

# Final note
fin = doc.add_paragraph(); fin.alignment = WD_ALIGN_PARAGRAPH.CENTER
fin.paragraph_format.space_before = Pt(20)
r = fin.add_run(f"Documento generado el {fecha_str} · Presupuesto N° {nro} · Válido por 30 días")
r.font.size = Pt(9); r.font.italic = True; r.font.color.rgb = SLATE_500

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUT))
print(f"OK - Presupuesto generado: {OUT}")
print(f"Tamaño: {OUT.stat().st_size:,} bytes")
