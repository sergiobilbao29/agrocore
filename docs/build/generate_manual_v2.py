"""
Manual de Usuario AgroCore v1.2 con imágenes ilustrativas.
Incluye todo el contenido del manual web actualizado.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

BASE = Path("/sessions/amazing-trusting-einstein/mnt/AgroCore")
IMGS = BASE / "docs/build/imgs"
LOGO_FULL = BASE / "web/img/logo-full-512.png"
OUT = BASE / "docs/AgroCore - Manual de Usuario.docx"

AGRO_DARK = RGBColor(0x14, 0x53, 0x2D)
AGRO_GREEN = RGBColor(0x15, 0x80, 0x3D)
AGRO_MED = RGBColor(0x16, 0x65, 0x34)
GOLD = RGBColor(0xCA, 0x8A, 0x04)
GOLD_LIGHT = RGBColor(0xEA, 0xB3, 0x08)
SLATE_700 = RGBColor(0x33, 0x41, 0x55)
SLATE_500 = RGBColor(0x64, 0x74, 0x8B)
AMBER_DARK = RGBColor(0xB4, 0x53, 0x09)

def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)

def set_cell_borders(cell, color="CCCCCC", size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), size); b.set(qn('w:color'), color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)

def add_page_number(paragraph):
    run = paragraph.add_run()
    f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = 'PAGE'
    f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'end')
    run._r.append(f1); run._r.append(it); run._r.append(f2)

def add_horizontal_line(paragraph, color="22C55E"):
    p_pr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '12'); bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), color)
    pbdr.append(bottom)
    p_pr.append(pbdr)

def h1(doc, text, color=AGRO_DARK):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16); p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name='Calibri'; r.font.size=Pt(20); r.font.bold=True; r.font.color.rgb=color
    add_horizontal_line(p, "DCFCE7")

def h2(doc, text, color=AGRO_GREEN):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name='Calibri'; r.font.size=Pt(15); r.font.bold=True; r.font.color.rgb=color

def h3(doc, text, color=AGRO_MED, badge=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name='Calibri'; r.font.size=Pt(12.5); r.font.bold=True; r.font.color.rgb=color
    if badge:
        r2 = p.add_run("  " + badge)
        r2.font.name='Calibri'; r2.font.size=Pt(9); r2.font.bold=True; r2.font.color.rgb=GOLD

def para(doc, text, bold=False, italic=False, color=None, size=11, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if align: p.alignment = align
    r = p.add_run(text)
    r.font.name='Calibri'; r.font.size=Pt(size); r.font.bold=bold; r.font.italic=italic
    if color: r.font.color.rgb = color
    return p

def para_mixed(doc, parts, size=11, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if align: p.alignment = align
    for text, opts in parts:
        r = p.add_run(text)
        r.font.name = 'Consolas' if opts.get('code') else 'Calibri'
        r.font.size = Pt(size - 1 if opts.get('code') else size)
        r.font.bold = opts.get('bold', False); r.font.italic = opts.get('italic', False)
        if opts.get('color'): r.font.color.rgb = opts['color']
    return p

def bullet(doc, text_or_parts):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.3)
    if isinstance(text_or_parts, str):
        r = p.add_run(text_or_parts); r.font.name='Calibri'; r.font.size=Pt(11)
    else:
        for text, opts in text_or_parts:
            r = p.add_run(text)
            r.font.name = 'Consolas' if opts.get('code') else 'Calibri'
            r.font.size = Pt(10 if opts.get('code') else 11)
            r.font.bold = opts.get('bold', False)
            if opts.get('color'): r.font.color.rgb = opts['color']

def numbered(doc, text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(2); p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(text); r.font.name='Calibri'; r.font.size=Pt(11)

def callout(doc, text, bg="F0FDF4", border="BBF7D0", fg=AGRO_GREEN, icon="💡"):
    table = doc.add_table(rows=1, cols=1); table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0); cell.width = Inches(6.5)
    set_cell_shading(cell, bg); set_cell_borders(cell, border, "4")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"{icon}  "); r.font.size=Pt(11)
    r = p.add_run(text); r.font.size=Pt(10.5); r.font.italic=True; r.font.color.rgb=fg
    after = doc.add_paragraph(); after.paragraph_format.space_after = Pt(0)
    return table

def section_banner(doc, image_filename, width_in=6.5):
    """Inserta una imagen de banner ancha como header de sección."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(8)
    p.add_run().add_picture(str(IMGS / image_filename), width=Inches(width_in))

def figure(doc, image_filename, caption=None, width_in=6.0):
    """Inserta una figura con caption opcional."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(IMGS / image_filename), width=Inches(width_in))
    if caption:
        cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(10)
        r = cp.add_run(caption); r.font.size=Pt(9); r.font.italic=True; r.font.color.rgb=SLATE_500

def styled_table(doc, headers, rows, col_widths_in=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers)); table.autofit = False
    if col_widths_in:
        for i, w in enumerate(col_widths_in): table.columns[i].width = Inches(w)
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        c = hdr.cells[i]
        if col_widths_in: c.width = Inches(col_widths_in[i])
        set_cell_shading(c, "166534"); set_cell_borders(c, "0F4626", "4")
        p = c.paragraphs[0]
        r = p.add_run(h); r.font.name='Calibri'; r.font.size=Pt(10.5); r.font.bold=True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for ri, row in enumerate(rows):
        tr = table.rows[ri+1]
        bg = "F8FAFC" if ri % 2 else "FFFFFF"
        for ci, val in enumerate(row):
            c = tr.cells[ci]
            if col_widths_in: c.width = Inches(col_widths_in[ci])
            set_cell_shading(c, bg); set_cell_borders(c, "E2E8F0", "4")
            p = c.paragraphs[0]
            r = p.add_run(str(val)); r.font.name='Calibri'; r.font.size=Pt(10)

# ============================================================
# BUILD
# ============================================================
doc = Document()

for section in doc.sections:
    section.page_height = Cm(29.7); section.page_width = Cm(21.0)
    section.top_margin = Cm(2.0); section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0); section.right_margin = Cm(2.0)

style = doc.styles['Normal']
style.font.name = 'Calibri'; style.font.size = Pt(11); style.font.color.rgb = SLATE_700

# Footer
footer = doc.sections[0].footer
fp = footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = fp.add_run("AgroCore Argentina · Manual de Usuario v1.2 · Página ")
r.font.name='Calibri'; r.font.size=Pt(9); r.font.color.rgb=SLATE_500
add_page_number(fp)

# Header
header = doc.sections[0].header
hp = header.paragraphs[0]; hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = hp.add_run("agrocore.ar")
r.font.name='Calibri'; r.font.size=Pt(9); r.font.color.rgb=AGRO_GREEN; r.font.bold=True

# ============================================================
# COVER
# ============================================================
# Cover decoration banner
cd = doc.add_paragraph(); cd.alignment = WD_ALIGN_PARAGRAPH.CENTER
cd.paragraph_format.space_before = Pt(0); cd.paragraph_format.space_after = Pt(20)
cd.add_run().add_picture(str(IMGS / "cover_decoration.png"), width=Inches(7))

cover_p = doc.add_paragraph(); cover_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover_p.paragraph_format.space_before = Pt(20)
if LOGO_FULL.exists():
    cover_p.add_run().add_picture(str(LOGO_FULL), width=Inches(3))

t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.paragraph_format.space_before = Pt(20)
r = t.add_run("Manual de Usuario")
r.font.name='Calibri'; r.font.size=Pt(36); r.font.bold=True; r.font.color.rgb=AGRO_DARK

sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("El corazón del negocio agrícola")
r.font.name='Calibri'; r.font.size=Pt(14); r.font.italic=True; r.font.color.rgb=GOLD

descr = doc.add_paragraph(); descr.alignment = WD_ALIGN_PARAGRAPH.CENTER
descr.paragraph_format.space_before = Pt(30)
r = descr.add_run("Sistema integral de gestión para empresas rurales y agropecuarias.")
r.font.name='Calibri'; r.font.size=Pt(12); r.font.color.rgb=SLATE_700
descr2 = doc.add_paragraph(); descr2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = descr2.add_run("Multi-empresa · Multi-usuario · En la nube o servidor propio.")
r.font.name='Calibri'; r.font.size=Pt(12); r.font.color.rgb=SLATE_700

ver = doc.add_paragraph(); ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
ver.paragraph_format.space_before = Pt(60)
r = ver.add_run("Versión 1.2 · Mayo 2026 · Sistema AgroCore v0.2")
r.font.name='Calibri'; r.font.size=Pt(10); r.font.color.rgb=SLATE_500
sop = doc.add_paragraph(); sop.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sop.add_run("Soporte vía WhatsApp / mail · Respuesta en 24 hs hábiles")
r.font.name='Calibri'; r.font.size=Pt(10); r.font.color.rgb=SLATE_500

doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

# ============================================================
# TOC
# ============================================================
h1(doc, "Tabla de contenidos")
toc_items = [
    "1. Bienvenida",
    "2. Acceder al sistema",
    "3. Primer login y conceptos clave",
    "4. Interfaz y navegación",
    "5. Trabajar con múltiples empresas",
    "6. Módulos del sistema",
    "    6.1 Inicio",
    "    6.2 Dashboard",
    "    6.3 Resumen multi-empresa  [NUEVO]",
    "    6.4 Producción",
    "    6.5 Stock",
    "    6.6 Contactos",
    "    6.7 Finanzas",
    "    6.8 Logística y RRHH",
    "    6.9 Administración",
    "7. Flujos de trabajo típicos",
    "8. Preguntas frecuentes",
    "9. Soporte y contacto",
]
for item in toc_items:
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2); p.paragraph_format.left_indent = Inches(0.2)
    r = p.add_run(item); r.font.name='Calibri'; r.font.size=Pt(11)
    if not item.startswith('    '):
        r.font.bold = True; r.font.color.rgb = AGRO_DARK
    else:
        r.font.color.rgb = SLATE_700
        if 'NUEVO' in item:
            r.font.color.rgb = GOLD

doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

# ============================================================
# 1. BIENVENIDA
# ============================================================
h1(doc, "1. Bienvenida")
para_mixed(doc, [
    ("AgroCore es un sistema integral pensado ", {}),
    ("por y para el agro argentino", {"bold": True}),
    (". No es un ERP genérico adaptado a la fuerza: cada módulo se diseñó pensando en las particularidades del campo argentino — campañas de gruesa y fina, arrendamientos en quintales, facturación con ARCA, libro IVA, cotizaciones de la BCR Rosario en vivo.", {}),
])
para(doc, "El sistema te permite gestionar:")
for label, desc in [
    ("Producción:", "campos, lotes, campañas, insumos y labores aplicadas."),
    ("Stock:", "granos y productos en general, con trazabilidad y reporte de margen."),
    ("Contactos:", "clientes y proveedores, con cuentas corrientes integradas."),
    ("Finanzas:", "flujo de caja, facturación electrónica ARCA, compras, libro IVA, cheques, cuentas corrientes, arrendamientos y control de efectivo."),
    ("Logística y RRHH:", "viajes (cartas de porte) y empleados (contratos, recibos)."),
    ("Administración:", "catálogos, usuarios, roles y permisos, empresas."),
]:
    bullet(doc, [(label + " ", {"bold": True}), (desc, {})])
bullet(doc, [("Resumen multi-empresa: ", {"bold": True, "color": GOLD}), ("vista consolidada de cheques, efectivo y flujo a través de todas las empresas a las que tengas acceso. ", {}), ("[NUEVO]", {"bold": True, "color": GOLD})])

para(doc, "Todo en una sola plataforma, con datos compartidos entre módulos. Si cargás una factura de venta, automáticamente se actualiza el stock, el flujo de caja proyectado, la cuenta corriente del cliente, el libro IVA y los KPIs del dashboard.")

# FIGURE: flow diagram
figure(doc, "flow_diagram.png", "Una sola operación actualiza varios módulos automáticamente.")

# ============================================================
# 2. ACCEDER AL SISTEMA
# ============================================================
h1(doc, "2. Acceder al sistema")
h2(doc, "Para el demo")
para(doc, "Durante el período de prueba, accedé desde cualquier navegador (Chrome, Edge, Firefox o Safari) en computadora, tablet o celular:")
para_mixed(doc, [("https://demo.agrocore.ar", {"code": True, "color": GOLD})])
para(doc, "El sistema te lleva automáticamente a la pantalla de login.")

h2(doc, "Para producción (cuenta propia)")
para(doc, "Una vez contratado, AgroCore se instala sobre tu propia infraestructura, a elección:")
bullet(doc, [("Cloud (AWS):", {"bold": True}), (" sobre tu cuenta de Amazon Web Services. URL personalizada (por ejemplo app.tuempresa.com.ar). Acceso desde cualquier lugar con internet.", {})])
bullet(doc, [("Servidor propio (on-premise):", {"bold": True}), (" instalado en un servidor físico de tu empresa. Acceso desde la red interna o vía VPN.", {})])
para_mixed(doc, [("En ambos casos, el sistema y los datos son ", {}), ("de tu propiedad", {"bold": True}), (".", {})])

h2(doc, "Compatibilidad")
bullet(doc, [("Navegadores:", {"bold": True}), (" Chrome, Edge, Firefox, Safari (últimos 2 años).", {})])
bullet(doc, [("Dispositivos:", {"bold": True}), (" computadora, notebook, tablet, celular.", {})])
bullet(doc, [("Sin instalación:", {"bold": True}), (" todo corre en el navegador.", {})])
bullet(doc, [("Modo offline:", {"bold": True}), (" funciona como PWA (Progressive Web App).", {})])

# ============================================================
# 3. LOGIN
# ============================================================
h1(doc, "3. Primer login y conceptos clave")
h2(doc, "Usuarios de prueba")
para(doc, "En el demo, ya hay dos usuarios cargados:")
styled_table(doc,
    ["Usuario", "Contraseña", "Rol"],
    [["Admin", "admin123", "Administrador de empresa"],
     ["Super", "super123", "Super administrador (todas las empresas)"]],
    col_widths_in=[1.3, 1.5, 3.7])
callout(doc, "En tu instalación productiva estos usuarios no existen — el primer alta la hace el equipo de implementación con tus datos reales.")

h2(doc, "Cómo iniciar sesión")
numbered(doc, "Entrá a https://demo.agrocore.ar")
numbered(doc, "Escribí el usuario (alias, nombre o email) y la contraseña.")
numbered(doc, "Clic en Ingresar.")

h2(doc, "Conceptos fundamentales")
h3(doc, "Empresa")
para(doc, "Cada empresa es una unidad independiente de datos: tiene sus campos, su stock, sus clientes, sus facturas. No se mezcla con otras. Una persona puede tener acceso a varias empresas (por ejemplo, un contador que lleva tres explotaciones).")
h3(doc, "Usuario")
para(doc, "La persona física que usa el sistema. Tiene un alias, email, foto opcional y contraseña encriptada.")
h3(doc, "Rol y permisos")
para_mixed(doc, [
    ("Cada usuario, en cada empresa donde tiene acceso, tiene un rol. Los permisos están agrupados por módulo (", {}),
    ("produccion:read", {"code": True}), (", ", {}),
    ("stock:read", {"code": True}), (", ", {}),
    ("finanzas:read", {"code": True}),
    (", etc.) y los roles vienen pre-cargados con permisos típicos del agro, personalizables por empresa.", {}),
])

# ============================================================
# 4. INTERFAZ Y NAVEGACIÓN
# ============================================================
h1(doc, "4. Interfaz y navegación")
h2(doc, "El layout general")
bullet(doc, [("Barra lateral (izquierda):", {"bold": True}), (" logo, selector de empresa activa, navegación entre módulos.", {})])
bullet(doc, [("Barra superior:", {"bold": True}), (" botón Volver, título de la pantalla, buscador global, notificaciones, ticker de cotizaciones BCR.", {})])
bullet(doc, [("Área principal (centro):", {"bold": True}), (" contenido del módulo actual.", {})])

h2(doc, "Las 7 secciones de la barra lateral")
styled_table(doc,
    ["Sección", "Módulos que contiene"],
    [
        ["INICIO", "Inicio · Dashboard · Resumen multi-empresa [NUEVO]"],
        ["PRODUCCIÓN", "Campos y lotes · Campañas · Insumos y labores"],
        ["STOCK", "Stock · Movimientos · Reporte de margen"],
        ["CONTACTOS", "Clientes · Proveedores"],
        ["FINANZAS", "Flujo de caja · Facturación · Compras · Libro IVA · Cheques · Cuentas corrientes · Arrendamientos · Control de efectivo"],
        ["LOGÍSTICA Y RRHH", "Viajes · Empleados"],
        ["ADMINISTRACIÓN", "Catálogos · Usuarios · Roles y permisos · Empresas"],
    ],
    col_widths_in=[1.8, 4.7])

h3(doc, "Botón Volver y navegación con historial", badge="[NUEVO]")
para_mixed(doc, [("En cualquier pantalla (excepto Inicio) verás un botón ", {}), ("← Volver", {"bold": True}), (" arriba a la izquierda del header. También funciona el botón Atrás del navegador o del celular: navega entre módulos dentro de la app, no te saca afuera.", {})])
para_mixed(doc, [("La URL refleja la página actual (ej. ", {}), ("/app#/cheques", {"code": True}), ("), así podés compartir links directos a vistas específicas — el destinatario abre el link y va directo a esa pantalla.", {})])

h2(doc, "El ticker de cotizaciones BCR")
styled_table(doc,
    ["Indicador", "Qué es"],
    [
        ["Blue", "Cotización del dólar paralelo (ARS por USD)"],
        ["CCL", "Contado con liquidación"],
        ["Cripto", "Promedio del USDT en exchanges argentinos"],
        ["Tarjeta", "Dólar oficial + impuestos"],
        ["Soja", "Precio Cámara Arbitral de Cereales BCR (USD/tonelada)"],
        ["Maíz", "Ídem soja"],
        ["Trigo", "Ídem soja"],
    ],
    col_widths_in=[1.3, 5.2])

h3(doc, "Notificaciones")
para_mixed(doc, [
    ("El icono de campanita arriba a la derecha muestra alertas: ", {}),
    ("cheques vencidos sin cobrar", {"bold": True}),
    (", ", {}),
    ("cheques a vencer en los próximos 15 días", {"bold": True, "color": GOLD}),
    (" (alerta proactiva ", {}),
    ("[NUEVO]", {"bold": True, "color": GOLD}),
    ("), stock bajo, facturas vencidas, vencimientos de arrendamientos.", {}),
])

h3(doc, "Buscador global")
para(doc, "La barra de búsqueda al centro-arriba permite encontrar rápidamente clientes, productos o empleados sin entrar al módulo correspondiente.")

# ============================================================
# 5. MULTI-EMPRESA
# ============================================================
h1(doc, "5. Trabajar con múltiples empresas")
para_mixed(doc, [
    ("Una funcionalidad clave: ", {}), ("gestionás varias empresas con el mismo usuario", {"bold": True}), (", sin cerrar sesión.", {}),
])
h2(doc, "Cómo cambiar de empresa")
numbered(doc, "Clic sobre el nombre de la empresa actual en la barra lateral.")
numbered(doc, "Aparece un menú con todas las empresas a las que tenés acceso.")
numbered(doc, "Clic en la empresa destino.")
numbered(doc, "La página se recarga mostrando los datos de la nueva empresa.")
para_mixed(doc, [("Al cambiar, ", {}), ("todo el sistema cambia de contexto", {"bold": True}), (". Sin riesgo de mezcla de datos.", {})])

h2(doc, "Color institucional por empresa")
para(doc, "Cada empresa puede tener configurado un color hex propio. Al cambiar de empresa, el borde activo del menú lateral y otros acentos visuales toman ese color. Útil cuando un contador lleva varias explotaciones — el color te recuerda en cuál estás trabajando.")

h2(doc, "Permisos por empresa")
para(doc, "Una persona puede ser Admin en la empresa A y Solo lectura en la empresa B. El sistema aplica el rol correspondiente al cambiar.")

h2(doc, "Para Super Admin")
para_mixed(doc, [("Los usuarios con rol ", {}), ("super", {"code": True}), (" pueden ver todas las empresas, crearlas/modificarlas/eliminarlas y asignar usuarios. Generalmente reservado para el dueño del sistema o el contador externo.", {})])

# ============================================================
# 6. MÓDULOS
# ============================================================
doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
h1(doc, "6. Módulos del sistema")

# 6.1 Inicio
section_banner(doc, "section_inicio.png")
h2(doc, "6.1 Inicio")
para(doc, "Punto de aterrizaje al entrar. Muestra:")
bullet(doc, [("Saludo personalizado", {"bold": True}), (" con tu nombre, empresa activa y fecha.", {})])
bullet(doc, [("KPIs rápidos", {"bold": True}), (" (lado derecho): Campañas activas, Stock granos (kg), Cabezas.", {})])
bullet(doc, [("Accesos rápidos", {"bold": True}), (": tarjetas con link directo a Campañas, Stock, Movimientos, Margen, Facturación, Compras, Clientes, Proveedores, Dashboard.", {})])
bullet(doc, [("Novedades de AgroCore", {"bold": True}), (": últimas mejoras y features publicados.", {})])
bullet(doc, [("Botón Ir al Dashboard", {"bold": True}), (" para profundizar en métricas.", {})])

# 6.2 Dashboard
section_banner(doc, "section_dashboard.png")
h2(doc, "6.2 Dashboard")
para(doc, "Vista de control para dueño / gerente. Tabs internas:")
bullet(doc, [("Resumen:", {"bold": True}), (" indicadores económicos generales del mes vs mes anterior.", {})])
bullet(doc, [("Finanzas:", {"bold": True}), (" ventas, compras, margen, cuentas corrientes, posición de cheques.", {})])
bullet(doc, [("Stock:", {"bold": True}), (" valorización, productos con menor rotación, alertas de mínimo.", {})])
bullet(doc, [("Producción:", {"bold": True}), (" campañas activas, hectáreas por cultivo, rinde estimado vs real.", {})])
bullet(doc, [("Evolución patrimonial:", {"bold": True}), (" cómo evolucionó el stock + cuentas a lo largo del tiempo.", {})])
para(doc, "Cada tab tiene sus propios gráficos y KPIs. Filtrable por rango de fechas y moneda (pesos / dólares).")

h3(doc, "Sección Alertas (panel derecho)")
para(doc, "Muestra alertas accionables con botón Ir → que lleva al módulo correspondiente:")
figure(doc, "alerts_card.png", "Panel de alertas del Dashboard con los 5 tipos de aviso, incluida la nueva alerta proactiva de 15 días.")

# 6.3 Resumen multi-empresa
section_banner(doc, "section_resumen.png")
h2(doc, "6.3 Resumen multi-empresa  [NUEVO]")
para_mixed(doc, [
    ("Vista global a través de TODAS las empresas a las que tu usuario tiene acceso. Mantiene la separación por empresa (no mezcla datos) pero ", {}),
    ("suma totales arriba", {"bold": True}),
    (" para tener \"el todo\" de un vistazo.", {}),
])

figure(doc, "multiempresa_card.png", "Ejemplo del Resumen multi-empresa con totales consolidados arriba y desglose por empresa abajo. La 3ra empresa está desmarcada y no suma al total.")

h3(doc, "Qué muestra")
bullet(doc, [("Cheques en cartera", {"bold": True}), (" (cantidad + monto consolidado de las empresas seleccionadas)", {})])
bullet(doc, [("Cheques a vencer en 15 días", {"bold": True}), (" (alerta proactiva consolidada)", {})])
bullet(doc, [("Cheques vencidos sin cobrar", {"bold": True})])
bullet(doc, [("Saldo de efectivo y flujo de caja", {"bold": True}), (" (consolidado)", {})])
bullet(doc, [("Desglose por empresa", {"bold": True}), (": tabla con una fila por empresa y sus totales individuales.", {})])

h3(doc, "Filtro de empresas")
para(doc, "Cada fila de la tabla tiene un checkbox. Al desmarcar una empresa, sus datos dejan de sumarse en los totales (pero la fila queda visible con opacity reducida para que sepas cuál está excluida). El checkbox del header te permite marcar/desmarcar todas las empresas con un clic.")
callout(doc, "Tu selección se guarda automáticamente — la próxima vez que entres a este módulo, el sistema recuerda qué empresas tenías marcadas.")

h3(doc, "Atajo a una empresa específica")
para_mixed(doc, [
    ("Cada fila tiene un botón ", {}), ("Cheques →", {"bold": True}),
    (" que cambia automáticamente a esa empresa y abre el módulo de Cheques directamente. Ideal para revisar el detalle de los cheques de una empresa puntual sin tener que cambiar manualmente desde el menú lateral.", {}),
])

h3(doc, "Importante: no mezcla datos")
para_mixed(doc, [
    ("Esta vista ", {}), ("no mezcla", {"bold": True}),
    (" datos entre empresas: cada cheque, efectivo y movimiento de flujo sigue asignado a la empresa donde se cargó. La separación por empresa se mantiene siempre. Solo ", {}),
    ("agrupa los totales", {"italic": True}), (" arriba para que puedas verlos en conjunto.", {}),
])

# 6.4 Producción
section_banner(doc, "section_produccion.png")
h2(doc, "6.4 Producción")
para(doc, "Sección que agrupa la gestión de campos, lotes, campañas e insumos.")

h3(doc, "Campos y lotes")
para_mixed(doc, [("Qué hace: ", {"bold": True}), ("permite cargar las parcelas con sus datos catastrales, superficie y subdivisiones (lotes).", {})])
para_mixed(doc, [("Datos típicos: ", {"bold": True}), ("nombre, localidad, superficie, tenencia (propio / arrendado / aparcería), lotes con su superficie y ambiente.", {})])
para_mixed(doc, [("Lotes y ambientes: ", {"bold": True}), ("cada lote puede tener un ambiente asociado (alto/medio/bajo, suelo arcilloso/franco/arenoso) usado en reportes de margen por ambiente. Si el campo es arrendado, se vincula al contrato en Finanzas → Arrendamientos.", {})])

h3(doc, "Campañas")
para_mixed(doc, [("Qué hace: ", {"bold": True}), ("ciclo productivo de un cultivo en uno o varios lotes durante una temporada. Núcleo de la planificación productiva.", {})])
para_mixed(doc, [("Datos: ", {"bold": True}), ("nombre, cultivo, tipo (gruesa/fina/doble), fechas siembra/cosecha previstas y reales, lotes incluidos, variedad, densidad, rinde estimado y real.", {})])
para_mixed(doc, [("Estados: ", {"bold": True}), ("Planificada → Sembrada → En curso → Cosechada → Histórica.", {})])
para_mixed(doc, [("Margen bruto: ", {"bold": True}), ("a medida que cargás insumos y al cosechar registrás rinde + precio, el sistema calcula automáticamente el margen por hectárea, total y vs planificación.", {})])

h3(doc, "Insumos y labores")
para_mixed(doc, [("Qué hace: ", {"bold": True}), ("registra cada aplicación de insumo (semilla, fertilizante, agroquímico) y cada labor (siembra, pulverización, cosecha) sobre un lote, con fecha, dosis, costo y operario.", {})])

para_mixed(doc, [("Hectáreas aplicadas custom ", {"bold": True}), ("[NUEVO]", {"bold": True, "color": GOLD}), (": en el formulario podés especificar la ", {}), ("cantidad real de hectáreas donde se aplicó", {"bold": True}), (". Por defecto el sistema autocompleta con el total del lote, pero es editable. Útil cuando aplicaste solo en una parte del lote (ej. el alrededor, una franja específica, o un sector).", {})])

figure(doc, "hectareas_illustration.png", "Aplicación total (50 ha) vs aplicación parcial (15 ha en el alrededor). El sistema calcula el costo real preciso según las hectáreas reales aplicadas.")

bullet(doc, "Si la cantidad es menor al total del lote, la tabla marca esa aplicación con un asterisco amarillo \"*\" indicando que fue parcial.")
bullet(doc, [("El costo total se calcula como ", {}), ("costoHa × hectáreasAplicadas", {"code": True}), (" (no como ", {}), ("costoHa × hectáreasTotalesDelLote", {"code": True}), ("), lo que da el costo real preciso.", {})])
bullet(doc, "El reporte de margen toma el costo real, no el estimado por el total del lote.")
para_mixed(doc, [("Integración con stock: ", {"bold": True}), ("cuando cargás un insumo aplicado, ", {}), ("automáticamente se descuenta del stock", {"bold": True}), (" del producto correspondiente. No hay que cargar dos veces.", {})])

# 6.5 Stock
section_banner(doc, "section_stock.png")
h2(doc, "6.5 Stock")
para(doc, "Maneja todo lo que se compra, almacena, consume o vende.")

h3(doc, "Stock (Productos y existencias)")
para(doc, "Listado del inventario actual: cada producto con cantidad disponible, valorizado a precio de última compra o promedio ponderado.")
para_mixed(doc, [("Tipos de productos: ", {"bold": True}), ("insumos agrícolas, granos, combustible, repuestos y consumibles, otros.", {})])

h3(doc, "Movimientos")
para(doc, "Cada vez que entra o sale algo del stock, se registra un movimiento. Trazabilidad completa.")
para_mixed(doc, [("Tipos: ", {"bold": True}), ("entrada por compra, entrada por cosecha, salida por venta, salida por consumo interno, ajuste de inventario, transferencia entre depósitos.", {})])

h3(doc, "Reporte de margen")
para(doc, "Análisis económico que cruza compras vs ventas calculando costo promedio, precio promedio, % de margen y ganancia. Filtros por producto, campaña, cliente/proveedor o período. Exportable a Excel.")

# 6.6 Contactos
section_banner(doc, "section_contactos.png")
h2(doc, "6.6 Contactos")
h3(doc, "Clientes")
para(doc, "Carga y gestión de clientes a quienes les vendés.")
para_mixed(doc, [("Datos: ", {"bold": True}), ("razón social, CUIT, condición IVA, domicilio, contacto, condiciones comerciales, observaciones internas.", {})])
para_mixed(doc, [("Filtros: ", {"bold": True}), ("por nombre, CUIT, email, localidad, y por ", {}), ("Condición de IVA", {"bold": True}), (" (Responsable Inscripto, Monotributista, Exento, Consumidor Final, No Inscripto). El filtro normaliza valores legacy automáticamente.", {})])
para_mixed(doc, [("Información derivada: ", {"bold": True}), ("saldo cuenta corriente, última factura, total facturado en el año, cheques pendientes.", {})])

h3(doc, "Proveedores")
para(doc, "Espejo de Clientes pero para vendedores de insumos, contratistas, transportistas, etc. Mismos campos. Información derivada de saldo a pagar y cheques entregados pendientes.")

# 6.7 Finanzas
section_banner(doc, "section_finanzas.png")
h2(doc, "6.7 Finanzas")
para(doc, "La sección más amplia del sistema. Agrupa toda la gestión económica y financiera.")

h3(doc, "Flujo de caja")
para(doc, "Proyección financiera consolidada a partir de cobros esperados, pagos comprometidos, vencimientos de cheques, pagos de arrendamientos y otros compromisos cargados en el sistema.")
para_mixed(doc, [("Salida: ", {"bold": True}), ("gráfico y tabla diaria/semanal/mensual con saldo proyectado a 30/60/90 días. Anticipa baches financieros antes de que ocurran.", {})])

h3(doc, "Facturación (ventas)")
para(doc, "Emite facturas electrónicas (A, B, C, M) directamente vía web service de ARCA, con CAE en el momento.")
para_mixed(doc, [("Cómo emitir: ", {"bold": True})])
numbered(doc, "Finanzas → Facturación → Nueva factura.")
numbered(doc, "Seleccionar cliente.")
numbered(doc, "El sistema sugiere el tipo de comprobante según condición IVA tuya/del cliente.")
numbered(doc, "Cargar ítems (producto, cantidad, precio, IVA). Si es producto del stock, se descuenta automáticamente.")
numbered(doc, "El dropdown muestra el stock disponible y avisa si querés vender más de lo que hay.")
numbered(doc, "Verificar total, agregar observaciones.")
numbered(doc, "Emitir → conexión a ARCA, CAE, PDF generado.")
numbered(doc, "Listo para enviar por email al cliente.")
para_mixed(doc, [("Tipos soportados: ", {"bold": True}), ("Facturas A/B/C/M, Notas de crédito y débito, Recibos, Remitos, Comprobantes de cobro.", {})])
para_mixed(doc, [("Modo informal: ", {"bold": True}), ("si tu empresa no factura por ARCA, el sistema genera comprobantes internos sin CAE pero con el mismo formato.", {})])

h3(doc, "Compras")
para_mixed(doc, [("Qué hace: ", {"bold": True}), ("registra las facturas que ", {}), ("te emiten", {"bold": True}), (" los proveedores. Si el ítem es producto del stock, ", {}), ("se incrementa automáticamente", {"bold": True}), (" al confirmar. Soporta importación masiva desde Excel o desde el portal Mis Comprobantes de ARCA.", {})])

h3(doc, "Libro IVA")
para(doc, "Consolida todas las operaciones gravadas del período (ventas + compras) discriminadas por alícuota, listo para presentación de DDJJ.")
bullet(doc, [("Libro IVA Ventas:", {"bold": True}), (" facturas emitidas del mes, con base imponible y débito fiscal por alícuota (10.5%, 21%, 27%, exento).", {})])
bullet(doc, [("Libro IVA Compras:", {"bold": True}), (" facturas recibidas del mes, con base y crédito fiscal por alícuota.", {})])
bullet(doc, [("Resumen mensual:", {"bold": True}), (" débito fiscal, crédito fiscal y saldo a pagar / favor.", {})])
para_mixed(doc, [("Exportación: ", {"bold": True}), ("archivos compatibles con AFIP/ARCA (CITI o Mis Comprobantes), Excel, PDF.", {})])

h3(doc, "Cheques")
para_mixed(doc, [("Estados: ", {"bold": True}), ("En cartera → Endosado / Depositado → Acreditado / Rechazado.", {})])
para_mixed(doc, [("Datos: ", {"bold": True}), ("número, banco, sucursal, importe, fechas (emisión y pago), cliente/proveedor, cuenta corriente afectada, CUIT del librador.", {})])
para_mixed(doc, [("Alertas proactivas ", {"bold": True}), ("[NUEVO]", {"bold": True, "color": GOLD}), (": el sistema avisa cuando un cheque está próximo a vencer en los ", {}), ("próximos 15 días", {"bold": True}), (" — no esperás a que venza, lo gestionás antes. Aparece tanto en el Dashboard como en la campanita de notificaciones.", {})])

h3(doc, "Cuentas corrientes")
para(doc, "Para cada cliente y proveedor, débitos y créditos con saldo actualizado.")
para_mixed(doc, [("Conciliación: ", {"bold": True}), ("permite atar un cobro a una factura específica, dejando claro qué está pago.", {})])

h3(doc, "Arrendamientos")
para_mixed(doc, [("Qué hace: ", {"bold": True}), ("registra los contratos de arrendamiento de campos cuando explotás campos que no son propios.", {})])
para_mixed(doc, [("Datos por contrato: ", {"bold": True})])
bullet(doc, "Arrendador (CUIT, datos de contacto)")
bullet(doc, "Campo arrendado (vinculado al módulo Producción → Campos)")
bullet(doc, "Plazo del contrato (fecha inicio, fecha fin)")
bullet(doc, "Modalidad de pago: pesos, dólares, quintales fijos de soja/maíz, % del rinde")
bullet(doc, "Cantidad acordada")
bullet(doc, "Vencimientos (uno o varios pagos a lo largo del contrato)")
para_mixed(doc, [("Integración: ", {"bold": True}), ("los vencimientos aparecen automáticamente en el flujo de caja como pagos comprometidos. Si el contrato es en quintales, el sistema valoriza la cuota al precio BCR del día.", {})])
para_mixed(doc, [("Alertas: ", {"bold": True}), ("vencimientos próximos, contratos por terminar, renovaciones a evaluar.", {})])

h3(doc, "Control de efectivo")
para(doc, "Registro de movimientos de caja chica: ingresos/egresos en efectivo con saldo actualizado. Datos: fecha, tipo, concepto (sueldos, viáticos, gastos varios, retiro de socio), importe, comprobante asociado.")

# 6.8 Logística y RRHH
section_banner(doc, "section_logistica.png")
h2(doc, "6.8 Logística y RRHH")

h3(doc, "Viajes")
para(doc, "Registra los viajes de transporte cuando se trasladan granos a destino (acopio, exportadora, planta de proceso) o entre campos propios.")
para_mixed(doc, [("Datos: ", {"bold": True}), ("fecha salida/llegada, origen y destino, transportista, chofer, patente, tipo de carga, tonelaje cargado y pesado en destino, cartas de porte (CTG), costo del flete.", {})])
para_mixed(doc, [("Integración: ", {"bold": True}), ("los kilos pesados en destino pueden generar el movimiento de stock asociado.", {})])

h3(doc, "Empleados")
para(doc, "Alta y gestión del personal de la empresa: datos personales, contratos, sueldos, recibos.")
para_mixed(doc, [("Contratos: ", {"bold": True}), ("cada empleado puede tener uno o varios contratos a lo largo del tiempo (ej. por temporada). Histórico mantenido.", {})])
para_mixed(doc, [("Recibos de sueldo: ", {"bold": True}), ("generación mensual con cálculo automático de descuentos, aguinaldo, vacaciones, horas extras. Exportables a PDF.", {})])

# 6.9 Administración
section_banner(doc, "section_admin.png")
h2(doc, "6.9 Administración")

h3(doc, "Catálogos")
para(doc, "Master data que alimenta a todos los módulos: productos, cultivos, bancos, conceptos de tesorería, provincias y localidades. Sidebar con filtros jerárquicos para encontrar rápido.")

h3(doc, "Usuarios")
para(doc, "Alta de personas que pueden iniciar sesión. Datos: email, alias, nombre/apellido, foto opcional, estado activo/inactivo, rol Super Admin opcional. Asignación a empresas con rol respectivo. Reseteo de contraseña por Admin.")

h3(doc, "Roles y permisos")
para(doc, "Roles pre-cargados (Admin, Contable, Operaciones, Lectura) o personalizados.")
bullet(doc, [("produccion:read / produccion:write", {"code": True})])
bullet(doc, [("stock:read / stock:write", {"code": True})])
bullet(doc, [("contactos:read / contactos:write", {"code": True})])
bullet(doc, [("finanzas:read / finanzas:write", {"code": True})])
bullet(doc, [("dashboard:read", {"code": True})])
bullet(doc, [("Wildcards: produccion:*, *:read, *:*", {"code": True})])

h3(doc, "Empresas")
para(doc, "Solo accesible para Super Admin. Datos: nombre, fantasía, CUIT, razón social, condición IVA, domicilio, logo, color institucional (hex), modo (con ARCA o informal), estado activa/inactiva. Operaciones: crear, editar, asignar usuarios y roles, configurar integración ARCA, desactivar.")

# ============================================================
# 7. FLUJOS DE TRABAJO
# ============================================================
doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
h1(doc, "7. Flujos de trabajo típicos")

h2(doc, "Flujo 1 — Cargar una compra de fertilizante")
numbered(doc, "Contactos → Proveedores → Nuevo (si no existe). CUIT, razón social, condición IVA.")
numbered(doc, "Finanzas → Compras → Nueva factura. Proveedor, tipo A, número, fecha, ítem (Urea Granulada × 5000 kg).")
numbered(doc, "Forma de pago: cuenta corriente. Confirmás.")
para_mixed(doc, [("Resultado automático: ", {"bold": True}), ("stock de Urea sube 5000 kg, cuenta corriente del proveedor con saldo deudor, aparece en flujo de caja, IVA registrado en libro IVA Compras.", {})])

h2(doc, "Flujo 2 — Aplicar el fertilizante a una campaña (con hectáreas parciales)")
numbered(doc, "Producción → Insumos y labores → + Insumo.")
numbered(doc, "Campaña (Soja 2025/26), lote (ej. lote de 80 ha), fecha.")
numbered(doc, "Insumo: Urea Granulada. Dosis: 100 kg/ha. Hectáreas aplicadas: 25 (solo el alrededor del lote, no completo).")
numbered(doc, "Sistema calcula: total = 2500 kg. Confirmás.")
para_mixed(doc, [("Resultado: ", {"bold": True}), ("stock de Urea baja 2500 kg, costo imputado a la campaña como aplicación parcial (marcado con * en la tabla).", {})])

h2(doc, "Flujo 3 — Vender granos y cobrar con cheques")
numbered(doc, "Finanzas → Facturación → Nueva factura. Cliente Acopiadora Pampeana, factura A, ítem Soja Cosecha 25/26 × 100 t × USD 380/t.")
numbered(doc, "Sistema emite, obtiene CAE, descuenta stock, genera PDF.")
numbered(doc, "Recibís el cheque. Finanzas → Cheques → Nuevo (recibido). Asociado a la factura, fecha de pago a 60 días.")
para_mixed(doc, [("Resultado: ", {"bold": True}), ("cuenta corriente del cliente al día, cheque en cartera, flujo de caja proyectado, sumado al libro IVA Ventas.", {})])
para_mixed(doc, [("15 días antes del vencimiento: ", {"bold": True, "color": GOLD}), ("el sistema lo muestra en la alerta \"⏰ Cheques a vencer en 15 días\" del Dashboard.", {})])

h2(doc, "Flujo 4 — Revisar consolidado de tres empresas en simultáneo  [NUEVO]")
numbered(doc, "Inicio → Resumen multi-empresa.")
numbered(doc, "Veo el total consolidado: cheques en cartera, a vencer en 15 días, vencidos, efectivo, flujo de caja.")
numbered(doc, "Si quiero excluir una empresa (ej. está en pausa), desmarco su checkbox → los totales se recalculan al instante.")
numbered(doc, "Una empresa muestra 3 cheques vencidos. Hago clic en Cheques → de esa fila.")
numbered(doc, "Sistema cambia automáticamente a esa empresa y abre Cheques, donde puedo ver el detalle y gestionarlos.")

h2(doc, "Flujo 5 — Cierre de campaña con rinde real")
numbered(doc, "Producción → Campañas → Editar Soja 2025/26. Cargás rinde real (35 qq/ha).")
numbered(doc, "Stock → Movimientos → Entrada por cosecha: 175 t entran al stock.")
numbered(doc, "Marcás campaña como Cosechada.")
para_mixed(doc, [("Resultado: ", {"bold": True}), ("stock sube, reporte de margen completo (ingresos - insumos - labores - arrendamiento si aplica).", {})])

h2(doc, "Flujo 6 — Pago de cuota de arrendamiento")
numbered(doc, "Finanzas → Arrendamientos → Editar contrato existente.")
numbered(doc, "Verificás el próximo vencimiento (ej. 200 quintales de soja, vence 30/06/2026).")
numbered(doc, "Al llegar el vencimiento, registrás el pago: el sistema valoriza al precio BCR del día.")
para_mixed(doc, [("Resultado: ", {"bold": True}), ("cuenta corriente del arrendador actualizada, libro IVA si corresponde, egreso en flujo de caja.", {})])

h2(doc, "Flujo 7 — Despacho de soja a destino")
numbered(doc, "Logística y RRHH → Viajes → Nuevo viaje.")
numbered(doc, "Origen: tu silo. Destino: Acopiadora Pampeana. Transportista, chofer, patente.")
numbered(doc, "Tonelaje cargado en origen, carta de porte (CTG).")
numbered(doc, "Al llegar a destino, cargás peso real y eventuales mermas.")
para_mixed(doc, [("Resultado: ", {"bold": True}), ("salida de stock, costo del flete imputado, asociación con factura de venta.", {})])

# ============================================================
# 8. FAQ
# ============================================================
doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
h1(doc, "8. Preguntas frecuentes")

faqs = [
    ("¿Pierdo datos si se corta la luz / internet?",
     "No. Los datos viven en el servidor (cloud o local), no en tu navegador. Si se corta la luz, los reanudás cuando vuelva. Si se corta internet pero el server local sigue prendido, podés seguir trabajando desde la red local. El modo PWA además guarda en el dispositivo lo que cargues offline."),
    ("¿Cómo se hacen los backups?",
     "Si estás en Cloud (AWS), automático y diario, con retención de 30 días. Si estás en servidor propio, según el plan acordado (típicamente diario incremental + semanal completo)."),
    ("¿Puedo migrar datos de mi sistema actual?",
     "Sí. Importamos desde Excel, planillas o archivos exportados. La migración inicial está incluida en la implementación."),
    ("¿Cuánta gente puede usar el sistema al mismo tiempo?",
     "No hay límite por licencia. Las únicas limitaciones son técnicas (servidor) y los precios planos no escalan por usuario."),
    ("¿Funciona desde el celular?",
     "Sí. Responsive y adaptado a pantallas pequeñas. La carga de movimientos puede hacerse desde el celular en el campo. El botón Atrás del celular navega entre módulos dentro de la app, no te saca afuera."),
    ("¿Cómo se factura con ARCA?",
     "Web service integrado. Una vez cargado tu certificado digital, las facturas se emiten directo con CAE."),
    ("¿Puedo ver datos de varias empresas en simultáneo?",
     "Sí, con el módulo Resumen multi-empresa (sección INICIO de la barra lateral). Sumás cheques, efectivo y flujo de caja de todas las empresas a las que tu usuario tiene acceso, sin perder la separación por empresa. Podés filtrar con checkboxes qué empresas incluir en el total."),
    ("¿Y si quiero un módulo nuevo o un cambio específico?",
     "Los pedidos se evalúan caso por caso. Hay un canal en tu plan de mantenimiento mensual. Cambios mayores tienen presupuesto aparte."),
    ("¿Quién es dueño de los datos?",
     "Vos. El código fuente y la infraestructura (cuenta AWS o servidor propio) son tuyos. Sin lock-in, sin sorpresas."),
    ("¿Cuánto tarda la implementación?",
     "Entre 4 y 6 semanas desde la confirmación, dependiendo del volumen de datos a migrar."),
    ("¿Mis datos están seguros?",
     "Estándares bancarios: HTTPS/TLS 1.3, encriptación AES-256, hash bcrypt para contraseñas, backups con retención, auditoría completa, OWASP Top 10. Firmamos NDA."),
    ("¿Puedo ver Hacienda en algún lado?",
     "En esta versión, los KPIs del Inicio muestran Cabezas como dato consolidado. La gestión completa de Hacienda (categorías, pesadas, sanidad) está prevista para una próxima versión. Si es crítico para tu operatoria, lo conversamos en el alcance de implementación."),
]
for q, a in faqs:
    h3(doc, q); para(doc, a)

# ============================================================
# 9. SOPORTE
# ============================================================
h1(doc, "9. Soporte y contacto")
para(doc, "Durante la prueba del demo y en el plan de mantenimiento:")
bullet(doc, [("WhatsApp:", {"bold": True}), (" (te lo pasamos en la primera conversación)", {})])
bullet(doc, [("Email:", {"bold": True}), (" sergiodbilbao@gmail.com", {})])
bullet(doc, [("Web:", {"bold": True}), (" https://agrocore.ar", {})])
para_mixed(doc, [("Horario: ", {"bold": True}), ("lunes a viernes, 9:00 a 18:00 hs (Argentina). ", {}), ("Respuesta: ", {"bold": True}), ("24 hs hábiles para consultas, 4 hs para incidentes críticos.", {})])

h2(doc, "Para reportar un bug durante el testing")
numbered(doc, "Qué estabas haciendo cuando pasó (paso a paso).")
numbered(doc, "Qué esperabas que pasara.")
numbered(doc, "Qué pasó en realidad.")
numbered(doc, "Captura de pantalla del error.")
numbered(doc, "Tu usuario y aproximadamente la hora.")

foot_p = doc.add_paragraph(); foot_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
foot_p.paragraph_format.space_before = Pt(40)
r = foot_p.add_run("Manual elaborado por el equipo de AgroCore Argentina.\nÚltima actualización: mayo 2026 — Versión 1.2.\n© AgroCore Argentina · Todos los derechos reservados.")
r.font.name='Calibri'; r.font.size=Pt(9); r.font.color.rgb=SLATE_500; r.font.italic=True

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUT))
print(f"OK - Manual v1.2 generado: {OUT}")
print(f"Tamaño: {OUT.stat().st_size:,} bytes")
