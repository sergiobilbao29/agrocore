"""Regenera el presupuesto agregando cover decoration banner."""
import subprocess, shutil
from pathlib import Path

# Solo correr el script existente para verificar que funciona
result = subprocess.run(['python3', 'build/generate_presupuesto.py'],
                       cwd='/sessions/amazing-trusting-einstein/mnt/AgroCore/docs',
                       capture_output=True, text=True)
print(result.stdout)
print(result.stderr if result.returncode else '')

# Ahora abro el .docx y le inserto el banner al principio
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOCX = Path("/sessions/amazing-trusting-einstein/mnt/AgroCore/docs/Presupuesto-AgroCore.docx")
BANNER = Path("/sessions/amazing-trusting-einstein/mnt/AgroCore/docs/build/imgs/cover_decoration.png")

doc = Document(str(DOCX))
# Insertar el banner como el primer elemento
body = doc.element.body
first_p = doc.paragraphs[0]
# Crear nuevo paragraph y moverlo al inicio
new_p = doc.add_paragraph()
new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
new_p.paragraph_format.space_before = Pt(0)
new_p.paragraph_format.space_after = Pt(15)
new_p.add_run().add_picture(str(BANNER), width=Inches(7))
# Mover el nuevo paragraph al inicio del body
body.insert(0, new_p._element)
# El nuevo paragraph se agregó al final también; lo removemos del final
# Actually, add_paragraph appends. We moved it to position 0. Remove duplicate.
doc.save(str(DOCX))
print(f"OK - banner agregado al presupuesto")
